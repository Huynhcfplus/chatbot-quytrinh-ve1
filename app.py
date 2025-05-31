import streamlit as st
from docx import Document
from io import BytesIO
import random

st.set_page_config(page_title="Chatbot Viết Quy trình", layout="centered")
st.title("🤖 Chatbot Viết Quy trình / Quy chế")

st.markdown("### 📝 Nhập thông tin để tạo quy trình")

# Mẫu quy trình có sẵn
mau_quytrinh = {
    "Quy trình tuyển dụng": {
        "muc_tieu": "Đảm bảo tuyển chọn đúng người, đúng vị trí.",
        "pham_vi": "Áp dụng cho toàn bộ các phòng ban.",
        "doi_tuong": "Bộ phận nhân sự và các trưởng bộ phận liên quan.",
        "buoc_thuc_hien": [
            "Xác định nhu cầu tuyển dụng",
            "Phê duyệt kế hoạch tuyển dụng",
            "Đăng tin và tiếp nhận hồ sơ",
            "Sàng lọc và phỏng vấn",
            "Thông báo kết quả và thử việc"
        ],
        "can_cu_phap_ly": "Bộ luật Lao động, Quy chế nhân sự công ty."
    },
    "Quy trình mua sắm thiết bị": {
        "muc_tieu": "Đảm bảo việc mua sắm đúng quy định, minh bạch.",
        "pham_vi": "Áp dụng cho tất cả bộ phận có nhu cầu mua thiết bị.",
        "doi_tuong": "Phòng HCNS và các bộ phận đề xuất.",
        "buoc_thuc_hien": [
            "Tiếp nhận yêu cầu mua sắm",
            "Phê duyệt kế hoạch mua sắm",
            "Lựa chọn nhà cung cấp",
            "Ký hợp đồng và nhận hàng",
            "Nghiệm thu và thanh toán"
        ],
        "can_cu_phap_ly": "Luật Đấu thầu, Quy định mua sắm nội bộ."
    }
}

# Chọn mẫu có sẵn
mau_chon = st.selectbox("📂 Chọn mẫu quy trình có sẵn (tùy chọn)", ["-- Tạo mới --"] + list(mau_quytrinh.keys()))

# Hiển thị tham khảo mẫu (nếu có chọn)
if mau_chon != "-- Tạo mới --":
    with st.expander("👁️ Xem nội dung mẫu tham khảo"):
        st.markdown(f"**Mục tiêu:** {mau_quytrinh[mau_chon]['muc_tieu']}")
        st.markdown(f"**Phạm vi:** {mau_quytrinh[mau_chon]['pham_vi']}")
        st.markdown(f"**Đối tượng:** {mau_quytrinh[mau_chon]['doi_tuong']}")
        st.markdown("**Các bước thực hiện:**")
        for idx, b in enumerate(mau_quytrinh[mau_chon]["buoc_thuc_hien"], 1):
            st.markdown(f"- Bước {idx}: {b}")
        st.markdown(f"**Căn cứ pháp lý:** {mau_quytrinh[mau_chon].get('can_cu_phap_ly', '')}")

with st.form("form"):
    ten_quytrinh = st.text_input("Tên quy trình / quy chế", value=mau_chon if mau_chon != "-- Tạo mới --" else "")
    ma_tai_lieu = st.text_input("Mã tài liệu (tùy chọn)", value=f"QT-{random.randint(100,999)}")
    linh_vuc = st.text_input("Lĩnh vực áp dụng")

    muc_tieu = st.text_area("Mục tiêu", value=mau_quytrinh.get(mau_chon, {}).get("muc_tieu", ""))
    pham_vi = st.text_area("Phạm vi áp dụng", value=mau_quytrinh.get(mau_chon, {}).get("pham_vi", ""))
    doi_tuong = st.text_area("Đối tượng thực hiện", value=mau_quytrinh.get(mau_chon, {}).get("doi_tuong", ""))
    buoc_thuc_hien = st.text_area("Các bước thực hiện (mỗi bước 1 dòng)",
                                  value="\n".join(mau_quytrinh.get(mau_chon, {}).get("buoc_thuc_hien", [])))
    can_cu_phap_ly = st.text_area("Căn cứ pháp lý / quy định liên quan",
                                  value=mau_quytrinh.get(mau_chon, {}).get("can_cu_phap_ly", ""))

    submitted = st.form_submit_button("Tạo quy trình")

if submitted:
    doc = Document()
    doc.add_heading(ten_quytrinh, 0)
    doc.add_paragraph(f"Mã tài liệu: {ma_tai_lieu}")
    doc.add_paragraph(f"Lĩnh vực: {linh_vuc}")
    doc.add_paragraph(f"Mục tiêu:\n{muc_tieu}")
    doc.add_paragraph(f"Phạm vi áp dụng:\n{pham_vi}")
    doc.add_paragraph(f"Đối tượng thực hiện:\n{doi_tuong}")
    doc.add_paragraph(f"Căn cứ pháp lý:\n{can_cu_phap_ly}")

    doc.add_heading("Các bước thực hiện", level=1)
    steps = buoc_thuc_hien.strip().split("\n")
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"Bước {i}: {step}", style="List Number")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    st.success("✅ Đã tạo quy trình thành công!")
    st.download_button(
        label="📄 Tải file Word",
        data=file_stream,
        file_name=f"{ten_quytrinh.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
